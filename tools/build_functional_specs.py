from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = r"C:\Users\nicol\Documents\ECE\ING2\semestre 2\projet piscine\Projet_Piscine\docs\Specifications_Fonctionnelles_Mercato_Nova.docx"


COLORS = {
    "blue": "2E74B5",
    "dark_blue": "1F4D78",
    "light_blue": "E8EEF5",
    "light_gray": "F2F4F7",
    "border": "B7C4D6",
    "muted": "666666",
    "risk": "9B1C1C",
}


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, COLORS["blue"], 16, 8),
        ("Heading 2", 13, COLORS["blue"], 12, 6),
        ("Heading 3", 12, COLORS["dark_blue"], 8, 4),
    ]:
        st = styles[name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor.from_string(color)
        st.font.bold = True
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True


def add_title(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("SPECIFICATIONS FONCTIONNELLES")
    run.bold = True
    run.font.size = Pt(23)
    run.font.color.rgb = RGBColor(0, 0, 0)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("Mercato Nova - Marketplace JDM voitures et pieces")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(55, 55, 55)

    metadata = [
        ("Projet", "Mercato Nova"),
        ("Version", "1.0 - specification fonctionnelle cible"),
        ("Base d'analyse", "Code React/Vite, squelette PHP, schema MySQL et decisions produit fournies"),
        ("Public cible", "Equipe produit, design, developpement, QA et enseignants"),
        ("Statut", "Document de reference pour implementation finale"),
    ]
    add_key_value_table(doc, metadata, [1800, 7560])

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("Synthese executif. ")
    r.bold = True
    p.add_run(
        "Le code actuel correspond a une maquette frontend avancee avec donnees simulees. "
        "La cible fonctionnelle attendue est une plateforme complete connectee a un backend PHP/MySQL, "
        "avec authentification, moderation admin/directeur, transactions par achat direct, enchere et negociation simple."
    )


def add_key_value_table(doc, rows, widths):
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
        set_cell_shading(cells[0], COLORS["light_gray"])
        for paragraph in cells[0].paragraphs:
            for run in paragraph.runs:
                run.bold = True
        for cell in cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
    set_table_width(table, widths)
    doc.add_paragraph()
    return table


def add_matrix_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, COLORS["light_blue"])
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.space_after = Pt(0)
            for run in paragraph.runs:
                run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
            for paragraph in cells[i].paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
    set_table_width(table, widths)
    doc.add_paragraph()
    return table


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_feature_detail(doc, name, objective, description, rules, preconditions, trigger, steps, result, messages, acceptance, code):
    doc.add_heading(name, level=3)
    add_key_value_table(
        doc,
        [
            ("Objectif", objective),
            ("Description fonctionnelle", description),
            ("Preconditions", preconditions),
            ("Declencheur", trigger),
            ("Resultat attendu", result),
            ("Fichiers / endpoints lies", code),
        ],
        [2100, 7260],
    )
    if rules:
        p = doc.add_paragraph()
        p.add_run("Regles metier.").bold = True
        add_bullets(doc, rules)
    if steps:
        p = doc.add_paragraph()
        p.add_run("Etapes du processus.").bold = True
        add_numbered(doc, steps)
    if messages:
        p = doc.add_paragraph()
        p.add_run("Messages utilisateur.").bold = True
        add_bullets(doc, messages)
    if acceptance:
        p = doc.add_paragraph()
        p.add_run("Criteres d'acceptation.").bold = True
        add_bullets(doc, acceptance)


def build_document():
    doc = Document()
    style_doc(doc)
    add_title(doc)

    doc.add_heading("1. Presentation du projet", level=1)
    add_key_value_table(
        doc,
        [
            ("Nom du projet", "Mercato Nova"),
            ("Contexte", "Projet piscine ECE ING2 visant a construire une marketplace automobile specialisee JDM."),
            ("Objectifs principaux", "Permettre la consultation d'offres, l'achat direct, les encheres, la negociation simple, la vente d'annonces validees, et l'administration de la plateforme."),
            ("Probleme a resoudre", "Centraliser des transactions automobiles avec controle de qualite des annonces et validation des vendeurs pour limiter les annonces frauduleuses."),
            ("Utilisateurs cibles", "Visiteurs, utilisateurs connectes acheteurs/vendeurs, administrateurs, directeur."),
            ("Parties prenantes", "Gabin, Nicolas, Paul, Celestin, enseignants et utilisateurs finaux de la maquette."),
        ],
        [2100, 7260],
    )

    doc.add_heading("2. Perimetre fonctionnel", level=1)
    add_matrix_table(
        doc,
        ["Inclus", "Description"],
        [
            ("Catalogue", "Consultation, recherche, filtres par categorie, marque, favoris et type de transaction."),
            ("Authentification", "Connexion par compte avec quatre niveaux d'acces : invite, compte connecte, administrateur, directeur."),
            ("Transactions", "Achat direct, enchere avec temps et prix de depart, negociation simple sans contre-offre pour la premiere version."),
            ("Vente", "Tout compte connecte peut proposer une annonce, soumise a validation administrateur ou directeur."),
            ("Administration", "Validation des annonces, verification carte d'identite uploadée, gestion des comptes, ban, suppression."),
            ("Direction", "Droits complets sur comptes, annonces et administrateurs, sans possibilite de modifier le prix d'une annonce."),
            ("Backend cible", "APIs PHP/MySQL pour remplacer les donnees mockees."),
        ],
        [2300, 7060],
    )
    add_matrix_table(
        doc,
        ["Exclu ou reporte", "Raison"],
        [
            ("Emails transactionnels", "Pas requis pour l'instant."),
            ("Negociation avancee avec contre-offres", "Explicitement exclue de la version actuelle."),
            ("Paiement reel", "Prevu a l'avenir via tunnel commande/paiement."),
            ("Validation automatique de carte d'identite", "La validation est manuelle par l'admin selon ses criteres."),
            ("Application mobile native", "Non visible dans le code ni demandee."),
        ],
        [2600, 6760],
    )
    doc.add_heading("Hypotheses", level=2)
    add_bullets(doc, [
        "Les quatre roles finaux sont : invite, utilisateur connecte, administrateur, directeur.",
        "Le role vendeur n'est pas separe : tout utilisateur connecte peut vendre apres validation de son annonce par un admin ou directeur.",
        "Le directeur dispose de tous les droits d'administration, mais ne peut pas modifier le prix d'une annonce, comme l'administrateur.",
        "La suppression d'un compte supprime aussi ses annonces, encheres et negociations associees.",
        "La base de donnees devra harmoniser ses enumerations avec les roles et types de vente de la cible fonctionnelle.",
    ])
    doc.add_heading("Contraintes", level=2)
    add_bullets(doc, [
        "Frontend actuel en React/Vite/Tailwind, backend cible en PHP pur, base MySQL.",
        "Les donnees frontend sont actuellement mockees ; aucune action metier n'est persistante en base.",
        "Les documents d'identite sont des donnees sensibles : stockage, acces et suppression devront etre securises.",
        "Le projet doit rester compatible avec un environnement local MAMP sous Windows.",
    ])

    doc.add_heading("3. Fonctionnalites existantes identifiees dans le code", level=1)
    add_matrix_table(
        doc,
        ["Fonctionnalite", "Description", "Modules", "Etat"],
        [
            ("Catalogue", "Affichage de six produits JDM avec image, prix, specs et statut.", "Home.jsx", "Partiel, mocke"),
            ("Recherche/filtres", "Filtrage par texte, categorie, marque, favoris et type de vente.", "Home.jsx", "Fonctionnel frontend"),
            ("Favoris", "Ajout/retrait de favoris en memoire React.", "Home.jsx", "Partiel, non persistant"),
            ("Panier", "Ajout, retrait, total et validation simulee.", "Home.jsx", "Partiel, mocke"),
            ("Enchere", "Modal, montant superieur au prix actuel, historique local.", "Home.jsx", "Partiel, mocke"),
            ("Negociation", "Ouverture simulee via alerte.", "Home.jsx", "Tres partiel"),
            ("Connexion", "Faux comptes, role selectionne, session localStorage.", "App.jsx, Login.jsx, mockData.js", "Partiel, mocke"),
            ("Admin", "Stats, annonces en attente, gestion comptes, ajout/modification/suppression annonces.", "Home.jsx, mockData.js", "Partiel, mocke"),
            ("Directeur", "Vue controle total annonces et acces comptes elargi.", "Home.jsx", "Partiel, mocke"),
            ("Backend", "Routeur racine et 404 JSON.", "backend/index.php", "Squelette"),
            ("Schema DB", "Users, products, auctions, negotiations.", "database/init.sql", "Defini, non connecte"),
        ],
        [1700, 3400, 2300, 1960],
    )

    doc.add_heading("4. Parcours utilisateurs", level=1)
    add_matrix_table(
        doc,
        ["Profil", "User stories principales"],
        [
            ("Invite", "En tant que visiteur, je consulte les offres et je filtre le catalogue afin d'evaluer les vehicules disponibles."),
            ("Utilisateur connecte", "En tant qu'utilisateur, je peux acheter, encherir, negocier et proposer une annonce afin de participer a la marketplace."),
            ("Administrateur", "En tant qu'administrateur, je valide les annonces et gere les comptes afin de proteger la qualite du site."),
            ("Directeur", "En tant que directeur, je supervise tous les comptes et toutes les annonces, y compris les administrateurs."),
        ],
        [2200, 7160],
    )
    doc.add_heading("Scenarios d'usage principaux", level=2)
    add_numbered(doc, [
        "Un invite arrive sur le catalogue, consulte les annonces, tente d'encherir, puis est redirige vers la connexion.",
        "Un utilisateur connecte ajoute une piece au panier et valide une commande simulee.",
        "Un utilisateur connecte propose une annonce ; elle apparait dans la file de moderation.",
        "Un administrateur examine une annonce, verifie l'identite du vendeur et accepte ou refuse la publication.",
        "Un administrateur bannit un compte ; la prochaine connexion de ce compte est refusee.",
        "Un directeur supprime un compte administrateur ; toutes les donnees liees doivent etre supprimees en base dans la cible.",
    ])
    doc.add_heading("Cas alternatifs et erreurs", level=2)
    add_bullets(doc, [
        "Identifiants invalides : afficher une erreur et rester sur la page connexion.",
        "Role selectionne incompatible avec le compte : refuser la connexion.",
        "Compte banni ou supprime : refuser la connexion.",
        "Offre d'enchere inferieure ou egale au prix actuel : refuser l'offre.",
        "Catalogue vide apres filtrage : afficher un etat vide avec action de reinitialisation.",
        "Action reservee a un utilisateur connecte : afficher un message puis ouvrir la connexion.",
    ])

    doc.add_heading("5. Description detaillee des fonctionnalites", level=1)
    add_feature_detail(
        doc,
        "Consultation et filtrage du catalogue",
        "Permettre a tous les visiteurs de consulter les offres disponibles.",
        "Le catalogue presente voitures et pieces avec informations techniques, image, prix, type de vente et statut de verification.",
        ["Le catalogue est accessible sans connexion.", "Les filtres peuvent etre combines.", "Un etat vide doit proposer une reinitialisation des filtres."],
        "Aucune connexion requise.",
        "Ouverture de l'application ou retour au catalogue.",
        ["Afficher les produits approuves.", "Appliquer recherche texte et filtres.", "Mettre a jour le nombre de resultats."],
        "L'utilisateur visualise uniquement les annonces publiees et approuvees.",
        ["Aucun bolide trouve dans la base.", "Reinitialiser les filtres."],
        ["Un invite peut afficher le catalogue.", "La recherche filtre marque, modele, moteur et chassis.", "Le filtre type de vente isole achat direct, enchere ou negociation."],
        "Home.jsx ; cible API GET /api/products",
    )
    add_feature_detail(
        doc,
        "Connexion, session et inscription",
        "Authentifier un utilisateur selon l'un des quatre roles retenus.",
        "La version actuelle utilise des comptes mockes ; la cible doit utiliser une authentification backend securisee.",
        ["Quatre roles : invite, utilisateur connecte, administrateur, directeur.", "Un compte banni ou supprime ne peut pas se connecter.", "L'inscription cible cree un compte utilisateur connecte standard."],
        "Utilisateur non connecte.",
        "Clic sur Se connecter ou action protegee.",
        ["Saisir email et mot de passe.", "Verifier le compte et son statut.", "Creer une session.", "Rediriger vers le catalogue ou l'espace demande."],
        "Session active avec role et permissions chargees.",
        ["Identifiants invalides.", "Role selectionne incompatible.", "Compte banni : connexion refusee."],
        ["Un compte valide se connecte.", "Un compte banni est refuse.", "Un utilisateur deconnecte redevient invite."],
        "App.jsx, Login.jsx, mockData.js ; cible POST /api/auth/login, POST /api/auth/register, POST /api/auth/logout",
    )
    add_feature_detail(
        doc,
        "Achat direct et panier",
        "Permettre l'achat d'une annonce en achat direct avec panier.",
        "Le panier actuel est local ; la cible prevoit un vrai tunnel de commande et paiement ulterieur.",
        ["La connexion est obligatoire pour acheter.", "Le panier affiche les articles, quantites implicites et total.", "Le paiement reel est reporte mais doit etre prevu dans l'architecture."],
        "Utilisateur connecte et produit disponible en achat direct.",
        "Clic sur Ajouter au panier puis Valider la commande.",
        ["Ajouter l'article.", "Afficher le compteur panier.", "Calculer le total.", "Lancer le tunnel de commande cible."],
        "Commande creee ou panier valide en simulation.",
        ["Connexion requise.", "Panier valide - simulation actuelle."],
        ["Un invite est redirige vers connexion.", "Un utilisateur connecte ajoute et retire un article.", "Le total se met a jour."],
        "Home.jsx ; cible POST /api/orders, POST /api/payments/init",
    )
    add_feature_detail(
        doc,
        "Encheres",
        "Permettre a un utilisateur connecte de placer une offre sur une annonce aux encheres.",
        "La cible comprend un temps d'enchere, un prix de depart et un gagnant final.",
        ["Connexion obligatoire.", "L'offre doit etre strictement superieure au prix courant.", "Une enchere possede une date de fin.", "A la fin, le plus haut encherisseur est gagnant."],
        "Produit de type enchere disponible.",
        "Clic sur Placer une offre.",
        ["Ouvrir la modal.", "Afficher prix actuel et historique.", "Saisir montant.", "Valider montant.", "Enregistrer l'offre.", "Mettre a jour le prix courant."],
        "Offre acceptee et nouveau prix courant affiche.",
        ["Votre offre doit etre strictement superieure a l'offre actuelle.", "Enchere enregistree."],
        ["Une offre trop basse est refusee.", "Une offre valide devient prix courant.", "La date de fin permet de designer un gagnant final."],
        "Home.jsx, database.auctions ; cible GET/POST /api/auctions, POST /api/auctions/{id}/bids",
    )
    add_feature_detail(
        doc,
        "Negociation simple",
        "Permettre d'ouvrir une intention de negociation sur une annonce eligible.",
        "La version cible n'inclut pas encore de contre-offres ou chat de negociation avance.",
        ["Connexion obligatoire.", "Pas de contre-offres pour l'instant.", "Une offre de negociation peut avoir un statut pending, accepted ou rejected en base."],
        "Utilisateur connecte et produit en negociation.",
        "Clic sur Negocier le prix.",
        ["Ouvrir le flux de negociation cible.", "Saisir une offre si le formulaire est implemente.", "Enregistrer l'offre."],
        "Negociation creee ou ouverture simulee selon l'etat d'implementation.",
        ["Connexion requise.", "Negociation ouverte."],
        ["Un invite ne peut pas negocier.", "Une negociation cree une ligne en base.", "Les statuts sont visibles dans l'historique cible."],
        "Home.jsx, database.negotiations ; cible POST /api/negotiations",
    )
    add_feature_detail(
        doc,
        "Depot et moderation d'annonce",
        "Permettre a tout utilisateur connecte de vendre, sous reserve de validation admin/directeur.",
        "L'annonce deposee n'est pas publiee directement ; elle rejoint une file de moderation avec carte d'identite uploadee.",
        ["Connexion obligatoire.", "Carte d'identite reelle uploadee.", "L'admin valide manuellement l'identite selon ses criteres.", "Une annonce refusee n'est pas publiee."],
        "Utilisateur connecte.",
        "Clic sur Proposer une annonce.",
        ["Saisir les informations annonce.", "Uploader carte d'identite.", "Envoyer en attente.", "Admin/directeur accepte ou refuse.", "Publier si acceptee."],
        "Annonce publiee apres validation, ou conservee en attente/refusee.",
        ["Annonce envoyee en moderation.", "Aucune annonce en attente."],
        ["Une annonce creee par utilisateur est invisible du catalogue tant qu'elle n'est pas approuvee.", "Admin/directeur peut accepter/refuser.", "L'acceptation ajoute l'annonce au catalogue."],
        "Home.jsx, mockData.js, database.products ; cible POST /api/products, GET /api/admin/listings/pending, POST /api/admin/listings/{id}/approve",
    )
    add_feature_detail(
        doc,
        "Gestion administrative des annonces",
        "Permettre aux administrateurs et directeurs de maintenir les annonces.",
        "Admin/directeur peuvent ajouter, modifier ou supprimer une annonce ; le prix est verrouille en modification.",
        ["Admin et directeur ont les memes restrictions sur le prix.", "Suppression d'une annonce retire aussi les encheres et negociations liees.", "Les modifications doivent etre journalisees dans la cible."],
        "Utilisateur avec role administrateur ou directeur.",
        "Clic sur croix de suppression, Modifier ou Ajouter une vente.",
        ["Ouvrir l'action.", "Modifier les champs autorises.", "Enregistrer.", "Mettre a jour le catalogue."],
        "Annonce mise a jour, creee ou supprimee.",
        ["Le prix ne peut pas etre modifie par un administrateur.", "Ajouter une vente."],
        ["Le prix reste identique apres modification.", "La suppression retire l'annonce du catalogue.", "L'ajout cree une annonce publiee ou soumise au statut choisi."],
        "Home.jsx ; cible PUT /api/admin/products/{id}, DELETE /api/admin/products/{id}",
    )
    add_feature_detail(
        doc,
        "Gestion des utilisateurs",
        "Permettre la supervision et la sanction des comptes.",
        "Admin et directeur consultent les comptes, bannissent ou suppriment. Le directeur peut agir aussi sur les administrateurs.",
        ["Admin peut voir les donnees personnelles de tous, y compris administrateurs.", "Directeur peut bannir/supprimer un administrateur.", "Suppression d'un compte supprime toutes les donnees liees.", "Donnees bancaires sensibles ne doivent pas etre exposees aux admins."],
        "Role administrateur ou directeur.",
        "Ouverture de Gerer les comptes.",
        ["Lister les comptes.", "Consulter les informations.", "Bannir ou supprimer.", "Appliquer les consequences sur la connexion et les donnees liees."],
        "Compte banni ou supprime selon action.",
        ["Compte banni : connexion refusee.", "Ce compte a ete supprime de la base de donnees simulee."],
        ["Un compte banni ne peut pas se reconnecter.", "Un compte supprime disparait de la liste.", "Les annonces du compte supprime sont supprimees en base cible."],
        "Home.jsx, mockData.js, database.users ; cible GET /api/admin/users, POST /api/admin/users/{id}/ban, DELETE /api/admin/users/{id}",
    )
    add_feature_detail(
        doc,
        "Notifications et messages in-app",
        "Informer l'utilisateur des evenements importants dans l'interface.",
        "La version actuelle affiche notifications et messages statiques ; aucun email n'est requis pour l'instant.",
        ["Notifications in-app uniquement pour la version cible initiale.", "Lecture d'une notification met son statut en lu.", "Messagerie actuelle simulee."],
        "Utilisateur sur l'application.",
        "Clic sur icone notification ou message.",
        ["Afficher la liste.", "Marquer une notification comme lue.", "Ouvrir une discussion ou message cible."],
        "Utilisateur voit les evenements et messages pertinents.",
        ["Tout lire.", "Discussion ouverte avec ..."],
        ["Le compteur non lu diminue apres lecture.", "Aucun email n'est envoye dans la premiere version."],
        "Home.jsx ; cible GET /api/notifications, PATCH /api/notifications/{id}/read",
    )

    doc.add_heading("6. Donnees et contenus", level=1)
    add_matrix_table(
        doc,
        ["Modele", "Champs principaux", "Regles de validation / remarques"],
        [
            ("User", "id, username, email, password_hash, role, first_name, last_name, phone, account_status, id_card_status, id_card_file, created_at", "Email unique, role parmi 4 niveaux fonctionnels, statut actif/banni/supprime."),
            ("Product", "id, brand, model, year, price, mileage, description, image_url, status, approval_status, sale_type, seller_id, approved_by", "Prix requis, prix non modifiable par admin/directeur apres creation, statut d'approbation obligatoire."),
            ("Auction", "id, product_id, start_price, current_bid, highest_bidder_id, start_date, end_date, status", "Date de fin requise, gagnant final = plus haute offre valide a cloture."),
            ("Bid", "id, auction_id, bidder_id, amount, created_at", "Montant strictement superieur au prix courant."),
            ("Negotiation", "id, product_id, buyer_id, offered_price, status, created_at", "Pas de contre-offre cible pour l'instant."),
            ("Order", "id, buyer_id, product_id, amount, status, payment_status, created_at", "A creer pour tunnel achat/paiement futur."),
        ],
        [1500, 4700, 3160],
    )
    doc.add_paragraph("Sources des donnees actuelles : tableaux React dans Home.jsx et mockData.js. Sources cibles : base MySQL via endpoints PHP.")

    doc.add_heading("7. Interfaces et ecrans", level=1)
    add_matrix_table(
        doc,
        ["Ecran", "Elements affiches", "Actions disponibles", "Etats attendus"],
        [
            ("Catalogue", "Navigation, recherche, filtres, cartes produit, favoris, panier.", "Filtrer, chercher, favori, acheter, encherir, negocier, proposer annonce.", "Chargement, vide, resultats, erreur API."),
            ("Connexion", "Email, mot de passe, role, bascule inscription, identifiants de test en maquette.", "Se connecter, s'inscrire, afficher mot de passe, retour catalogue.", "Erreur identifiants, compte banni, succes."),
            ("Panier", "Liste articles, total, bouton validation.", "Retirer article, valider commande.", "Vide, rempli, paiement futur."),
            ("Modal enchere", "Produit, prix actuel, champ montant, historique.", "Saisir montant, confirmer.", "Erreur montant, succes."),
            ("Console admin", "Stats, file moderation, gestion ventes, comptes.", "Valider/refuser, bannir, supprimer, modifier, ajouter.", "Vide, succes, erreur permissions."),
            ("Espace directeur", "Comptes incluant admins, annonces, controles globaux.", "Supprimer/bannir admin, supprimer/modifier annonces.", "Succes, erreur, confirmation suppression."),
        ],
        [1600, 2900, 2800, 2060],
    )

    doc.add_heading("8. Roles et permissions", level=1)
    add_matrix_table(
        doc,
        ["Action", "Invite", "Utilisateur connecte", "Administrateur", "Directeur"],
        [
            ("Consulter catalogue", "Oui", "Oui", "Oui", "Oui"),
            ("Acheter", "Non", "Oui", "Oui", "Oui"),
            ("Encherir", "Non", "Oui", "Oui", "Oui"),
            ("Negocier", "Non", "Oui", "Oui", "Oui"),
            ("Proposer annonce", "Non", "Oui, en attente validation", "Oui", "Oui"),
            ("Valider annonce", "Non", "Non", "Oui", "Oui"),
            ("Modifier annonce hors prix", "Non", "Non", "Oui", "Oui"),
            ("Modifier prix annonce", "Non", "Non", "Non", "Non"),
            ("Supprimer annonce", "Non", "Non", "Oui", "Oui"),
            ("Bannir/supprimer utilisateur", "Non", "Non", "Oui", "Oui"),
            ("Bannir/supprimer admin", "Non", "Non", "A confirmer si auto-protection seulement", "Oui"),
            ("Voir donnees sensibles bancaires", "Non", "Non", "Non", "A definir strictement"),
        ],
        [2500, 1600, 2200, 1500, 1560],
    )
    doc.add_paragraph("Logique observee : les permissions sont actuellement calculees cote frontend par role. La cible doit deplacer l'autorisation cote backend.")

    doc.add_heading("9. Notifications et communications", level=1)
    add_bullets(doc, [
        "Notifications in-app : depassement d'enchere, nouveau message, validation escrow ou moderation.",
        "Messages systeme : connexion requise, compte banni, annonce envoyee en moderation, enchere enregistree.",
        "Emails : exclus pour l'instant.",
        "Conditions d'envoi cibles : creation annonce, decision moderation, enchere depassee, enchere gagnee, commande creee.",
    ])

    doc.add_heading("10. Exigences non fonctionnelles", level=1)
    add_matrix_table(
        doc,
        ["Categorie", "Exigence"],
        [
            ("Performance", "Catalogue filtre localement en maquette ; cible paginee et indexee cote backend pour volumes importants."),
            ("Securite", "Authentification serveur, mots de passe hashes, sessions securisees, controles RBAC cote API."),
            ("Protection donnees", "Carte d'identite stockee de maniere securisee, acces limite, suppression lors de suppression du compte."),
            ("RGPD", "Droit de suppression applique aux comptes et donnees liees ; journalisation et conservation a preciser."),
            ("Accessibilite", "Boutons nommes, contrastes suffisants, navigation clavier a verifier."),
            ("Compatibilite", "Navigateurs modernes desktop/mobile, environnement local MAMP Windows."),
            ("Disponibilite", "Mode local projet ; cible non definie pour production."),
        ],
        [1800, 7560],
    )

    doc.add_heading("11. Integrations externes", level=1)
    add_matrix_table(
        doc,
        ["Integration", "Etat", "Details"],
        [
            ("MySQL", "Preparee", "Schema init.sql, connexion PDO database.php."),
            ("API PHP", "Squelette", "Routeur JSON avec endpoints annonces mais non implementes."),
            ("Images externes", "Utilisees", "URLs Unsplash et Googleusercontent pour la maquette."),
            ("Paiement", "A venir", "Tunnel commande/paiement a definir ; aucune integration actuelle."),
            ("Email", "Exclu V1", "Aucun envoi requis pour l'instant."),
            ("Upload fichiers", "A venir", "Dossier backend/uploads present ; carte d'identite reelle a uploader."),
        ],
        [1800, 1600, 5960],
    )
    doc.add_heading("Endpoints cibles recommandes", level=2)
    add_matrix_table(
        doc,
        ["Methode", "Endpoint", "Role requis", "Objectif"],
        [
            ("POST", "/api/auth/login", "Tous", "Connexion."),
            ("POST", "/api/auth/register", "Invite", "Creation compte utilisateur."),
            ("GET", "/api/products", "Tous", "Catalogue approuve."),
            ("POST", "/api/products", "Utilisateur+", "Depot annonce en attente avec upload CI."),
            ("PUT", "/api/admin/products/{id}", "Admin/Directeur", "Modification champs autorises hors prix."),
            ("DELETE", "/api/admin/products/{id}", "Admin/Directeur", "Suppression annonce et donnees liees."),
            ("GET", "/api/admin/listings/pending", "Admin/Directeur", "File moderation."),
            ("POST", "/api/admin/listings/{id}/approve", "Admin/Directeur", "Accepter annonce."),
            ("POST", "/api/admin/listings/{id}/reject", "Admin/Directeur", "Refuser annonce."),
            ("GET", "/api/admin/users", "Admin/Directeur", "Liste comptes."),
            ("POST", "/api/admin/users/{id}/ban", "Admin/Directeur", "Bannir compte."),
            ("DELETE", "/api/admin/users/{id}", "Admin/Directeur", "Supprimer compte et donnees liees."),
            ("POST", "/api/auctions/{id}/bids", "Utilisateur+", "Placer enchere."),
            ("POST", "/api/negotiations", "Utilisateur+", "Creer negociation simple."),
            ("POST", "/api/orders", "Utilisateur+", "Creer commande achat direct."),
        ],
        [900, 3100, 1700, 3660],
    )

    doc.add_heading("12. Criteres d'acceptation globaux", level=1)
    add_bullets(doc, [
        "Un invite peut consulter le catalogue mais ne peut pas acheter, vendre, encherir ou negocier sans connexion.",
        "Un utilisateur connecte peut acheter, encherir, negocier et deposer une annonce en attente.",
        "Une annonce deposee n'apparait pas au catalogue tant qu'un admin/directeur ne l'a pas acceptee.",
        "Admin et directeur peuvent modifier une annonce sans pouvoir changer son prix.",
        "Un directeur peut supprimer ou bannir un administrateur.",
        "La suppression d'un compte supprime toutes ses donnees liees : annonces, encheres, negociations, commandes et fichiers.",
        "Un compte banni ne peut pas se reconnecter.",
        "Les endpoints backend appliquent les permissions cote serveur, pas uniquement cote frontend.",
        "Le build frontend passe et les parcours principaux sont testables.",
        "Les tests QA couvrent les cas nominaux, alternatifs et erreurs listes dans ce document.",
    ])
    doc.add_heading("Tests fonctionnels attendus", level=2)
    add_matrix_table(
        doc,
        ["Parcours", "Validation attendue"],
        [
            ("Invite -> enchere", "Message connexion requise et ouverture page connexion."),
            ("Utilisateur -> depot annonce", "Annonce visible dans file moderation, pas dans catalogue public."),
            ("Admin -> acceptation annonce", "Annonce ajoutee au catalogue approuve."),
            ("Admin -> modification", "Tous les champs autorises changent, prix conserve."),
            ("Directeur -> suppression admin", "Compte admin et donnees liees supprimes."),
            ("Compte banni -> login", "Connexion refusee."),
            ("Enchere -> offre trop basse", "Erreur montant."),
            ("Enchere -> fin", "Gagnant final determine par plus haute offre."),
        ],
        [2800, 6560],
    )

    doc.add_heading("13. Questions ouvertes", level=1)
    add_bullets(doc, [
        "Quel prestataire de paiement sera choisi pour le tunnel futur ?",
        "Quelle duree de conservation appliquer aux cartes d'identite apres validation ou refus ?",
        "Faut-il journaliser les actions admin/directeur dans une table d'audit ?",
        "Quel format maximal accepter pour l'upload de carte d'identite : PDF, JPG, PNG, taille maximale ?",
        "Faut-il afficher un historique public des encheres ou seulement a l'admin/vendeur ?",
        "Faut-il une confirmation avant suppression definitive des comptes et annonces ?",
        "Le directeur peut-il consulter les donnees bancaires completes ou seulement des informations masquees et auditees ?",
        "La negociation simple doit-elle inclure un montant saisi des la V1 ou seulement une demande de contact ?",
    ])
    doc.add_heading("Risques identifies", level=2)
    add_bullets(doc, [
        "Incoherence actuelle entre roles frontend et roles SQL.",
        "Types de vente frontend et SQL a harmoniser pour supporter achat direct.",
        "Dependance actuelle a localStorage pour des actions sensibles.",
        "Dossier upload present mais aucune securite de fichier encore implementee.",
        "Regles de suppression en cascade a verifier soigneusement pour eviter des donnees orphelines.",
    ])

    doc.add_page_break()
    doc.add_heading("Annexe A - Cartographie technique actuelle", level=1)
    add_matrix_table(
        doc,
        ["Fichier", "Responsabilite actuelle"],
        [
            ("frontend/src/App.jsx", "Gestion affichage login/catalogue, session locale et rejet des comptes bannis/supprimes."),
            ("frontend/src/views/Login.jsx", "Formulaire connexion/inscription simulee et validation de faux comptes."),
            ("frontend/src/views/Home.jsx", "Catalogue, menus, panier, encheres, moderation, admin/directeur et modales."),
            ("frontend/src/data/mockData.js", "Comptes de test, annonces en attente, statistiques admin."),
            ("database/init.sql", "Schema MySQL et seed data pour users, products, auctions, negotiations."),
            ("backend/index.php", "Routeur API racine, CORS et 404."),
            ("backend/config/database.php", "Connexion PDO MySQL MAMP Windows."),
        ],
        [3000, 6360],
    )

    doc.core_properties.title = "Specifications Fonctionnelles - Mercato Nova"
    doc.core_properties.subject = "Marketplace JDM - specification fonctionnelle cible"
    doc.core_properties.author = "Codex"
    doc.save(OUTPUT)


if __name__ == "__main__":
    build_document()
